from fastapi import FastAPI
from fastapi.templating import Jinja2Templates
from starlette.responses import HTMLResponse
import sqlite3
from docx import Document
from update_database import get_all_applicant_data
import openai
from openai import OpenAI
import ast
import sqlite3
import re
import os


def create_CV(user_id, job_description, template_path):
    profile_dict = create_user_profile(user_id, job_description)
    update_word_document_with_user_info(template_path, profile_dict)
    return template_path.replace('.docx', '_updated.docx')


def create_user_profile(user_id, job_description):

    conn_DB = sqlite3.connect('database.db')

    profile = get_all_applicant_data(conn_DB, user_id)

    user_profile = {'Full name': profile['applicant']['full_name'].upper(),
                    'email place holder': profile['applicant']['email'],
                    'phone number place holder': profile['applicant']['phone_number'],
                    'About me place holder': profile['applicant']['professional_summary'],
                    'Exp place holder': profile['work_experience'],
                    'Education place holder': profile['education'],
                    'skills place holder': profile['skills'],
                    'vol place holder': profile['volunteering'],
                    'lan place holder': profile['languages']}

    client = None
    try:
        api_key = os.environ.get('API_KEY')
        client = OpenAI(api_key=api_key)
    except Exception as e:
        print(str(e))
    prompt = (f"""You are a professional recruiter, perfect for helping candidates get their dream jobs.
                     Given the following job description and a dictionary of applicant details, adjust, thicken   
                     and improve the dictionary such that it will highlight the applicants qualities needed for the job 
                     (without making anything up!!! it is very important that you won't add data that is not real 
                     about the candidate, and you can change only the values, leave the keys exactly the same).
                     Output only the python dictionary, without any explanation or details.
                     <applicant_data>: {user_profile}
                     </applicant_data>
                     <job_details>: {job_description}
                     </job_data>""")
    chat_completion = client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="gpt-3.5-turbo"
    )
    applicant_data = chat_completion.choices[0].message.content

    if isinstance(applicant_data, str):
        applicant_data = re.sub(r'([a-zA-Z0-9])\'([a-zA-Z0-9])', r'\1\2', applicant_data)

    applicant_data = ast.literal_eval(applicant_data)

    conn_DB.close()

    return applicant_data

def update_word_document_with_user_info(doc_path, user_info):
    doc = Document(doc_path)

    for paragraph in doc.paragraphs:
        update_paragraph(paragraph, user_info)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    update_paragraph(paragraph, user_info)

    new_doc_path = doc_path.replace('.docx', '_updated.docx')
    doc.save(new_doc_path)
    print(f"Document saved as '{new_doc_path}'.")

# def update_paragraph(paragraph, user_info):
#     for key, value in user_info.items():
#         for run in paragraph.runs:
#             if key.lower() in run.text.lower():
#                 # Replace key with value in text while preserving formatting
#                 run.text = run.text.lower().replace(key.lower(), value)

def update_paragraph(paragraph, user_info):
    """Check if any key in user_info is in the paragraph, and if so, replace it while preserving formatting."""
    for key, value in user_info.items():
        if isinstance(value, list):
            value_str = '\n'.join([
                '- ' + '. '.join([str(inner_item) for inner_item in item if str(inner_item).strip()]) + '.'
                if any(str(inner_item).strip() for inner_item in item) else ''  # Add a period at the end if the list is not empty
                for item in value
            ]).rstrip('\n')
        else:
            value_str = value

        for run in paragraph.runs:
            if key.lower() in run.text.lower():
                run.text = run.text.replace(key, value_str, 1)