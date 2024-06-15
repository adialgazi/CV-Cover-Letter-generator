import sqlite3
import json
from datetime import datetime
from docx import Document
from create_database import create_connection
from openai import OpenAI
import os


def fetch_applicant_data(uuid, conn):
    tables = [
        "applicants", "education", "projects", "work_experience",
        "volunteer_work", "languages", "certifications",
        "awards", "skills", "personal_projects"
    ]
    data = {}

    for table in tables:
        data[table] = []
        query = f"SELECT * FROM {table} WHERE applicant_uuid = ?"
        cursor = conn.cursor()
        cursor.execute(query, (uuid,))
        rows = cursor.fetchall()

        columns = [column[0] for column in cursor.description]

        for row in rows:
            row_data = dict(zip(columns, row))
            data[table].append(row_data)


    email = data['applicants'][0]['email']
    phone = data['applicants'][0]['phone_number']
    full_name = data['applicants'][0]['full_name']
    title = data['applicants'][0]['title']
    return data, email, phone, full_name, title

def generate_cover_letter(applicant_data, job_data):
    client = None
    try:
        api_key = os.environ.get('API_KEY')
        client = OpenAI(api_key=api_key)
    except Exception as e:
        print(str(e))
    prompt = (f"""You are a professional recruiter, perfect for helping candidates get their dream jobs.
                 Given the following job description and applicant details, write a cover letter that 
                 highlights the applicants qualities needed for the job (without making anything up).
                 The letter should be written in the 1st person, and should only take the relevant data to the job.
                 Output only the letter, without any explanation or details.
                 <applicant_data>: {applicant_data}
                 </applicant_data>
                 <job_details>: {job_data}
                 </job_data>""")
    chat_completion = client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="gpt-3.5-turbo"
    )
    ai_output = chat_completion.choices[0].message.content
    print(ai_output)
    return(ai_output)


def replace_placeholder_in_runs(runs, placeholder, replacement):
    if replacement is None:
        replacement = ''
    if placeholder == 'Phone Number':
        combined_text = ''.join(run.text for run in runs)
        placeholder_index = combined_text.find(placeholder)

        if placeholder_index == -1:
            return

        before_placeholder = combined_text[:placeholder_index]
        after_placeholder = combined_text[placeholder_index + len(placeholder):]

        new_combined_text = before_placeholder + replacement + after_placeholder

        current_index = 0
        for run in runs:
            run_text_length = len(run.text)
            run.text = new_combined_text[current_index:current_index + run_text_length]
            current_index += run_text_length

    else:
        combined_text = ''.join(run.text for run in runs)
        if placeholder in combined_text:
            start_index = combined_text.find(placeholder)
            end_index = start_index + len(placeholder)

            current_index = 0
            for run in runs:
                run_end_index = current_index + len(run.text)

                if start_index < run_end_index and end_index > current_index:
                    part_to_replace = combined_text[start_index:end_index]
                    run.text = run.text.replace(part_to_replace, replacement, 1)

                    offset = len(replacement) - len(part_to_replace)
                    start_index += offset
                    end_index += offset

                current_index += len(run.text)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        print(f"Found: {key}")
        replace_placeholder_in_runs(paragraph.runs, key, value)

def update_word_template(file_path, contact_tuple, cover_letter_text):
    doc = Document(file_path)

    today_date = datetime.today().strftime('%d %b, %Y')

    replacements = {
        "Title": contact_tuple[3],
        "Full Name": contact_tuple[2],  # Assuming index 2 is Full Name
        "Phone Number": contact_tuple[1],  # Assuming index 1 is Phone Number
        "Email": contact_tuple[0],  # Assuming index 0 is Email
        "Text": cover_letter_text,  # Cover letter text to replace 'Text' placeholder
        "Date": today_date  # Placeholder for today's date
    }

    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            replace_text_in_paragraph(paragraph, key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        replace_text_in_paragraph(paragraph, key, value)

    new_file_path = "temporary_files/modified_CL_template.docx"
    doc.save(new_file_path)
    return new_file_path


def main(job_details, uuid):
    database = "database.db"
    print('in main')
    conn = create_connection(database)
    print(conn)
    if conn is not None:
        applicant_data, email, phone, full_name, title = fetch_applicant_data(uuid, conn)
        print(applicant_data, email, phone, full_name, title)
        conn.close()
        data_tuple = (email, phone, full_name, title)
        applicant_data_json = json.dumps(applicant_data, indent=4)
        cover_letter_text = generate_cover_letter(applicant_data_json, job_details)

        updated_file_path = update_word_template('static/templates/CL_template.docx', data_tuple, cover_letter_text)

        print(f"Updated document saved to: {updated_file_path}")
        return updated_file_path
    else:
        print("Error! Cannot create the database connection.")

if __name__ == "__main__":
    job_details = "Your job details here"
    something = update_word_template('static/templates/CL_template.docx', ('my email', '123213', 'Ilai :)', 'Software guy'), 'I am cool')
